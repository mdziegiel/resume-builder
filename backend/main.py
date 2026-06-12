import io
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

import requests
from docx import Document
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from database import DB_PATH, UPLOAD_DIR, connect, dumps, loads, now
from document_renderer import docx_bytes, merge_resume_data, pdf_bytes, safe_filename
from seed import DEFAULT_RESUME, ROLE_LIBRARY, insert_sample_resume, seed

app = FastAPI(title='Resume Builder', version='1.0.0')
app.add_middleware(CORSMiddleware, allow_origins=['*'], allow_methods=['*'], allow_headers=['*'])

class ResumeIn(BaseModel):
    name: str
    title: str
    template: str = 'executive'
    data: dict[str, Any]

class ResumeRenderIn(BaseModel):
    template: str = 'modern'
    data: dict[str, Any]
    name: str = 'resume'

class DocIn(BaseModel):
    kind: str
    title: str
    resume_id: Optional[int] = None
    data: dict[str, Any]

class TailorIn(BaseModel):
    resume_id: int
    job_description: str

class CoverIn(BaseModel):
    resume_id: Optional[int] = None
    company: str
    job_description: str
    hiring_manager: str = 'Hiring Manager'
    sender_name: str = ''
    sender_contact: str = ''

class ThankYouIn(BaseModel):
    resume_id: Optional[int] = None
    interviewer: str
    company: str
    position: str
    interview_date: str
    talking_points: str
    sender_name: str = ''
    sender_contact: str = ''

class RoleBuildIn(BaseModel):
    title: str
    industry: str = 'Information Technology'
    category: Optional[str] = None

class ImportConfirmIn(BaseModel):
    filename: str = 'Imported Resume'
    template: str = 'modern'
    data: dict[str, Any]

class PersonalResumeIn(BaseModel):
    title: str
    years: str = ''
    skills: str = ''
    employer_role: str = ''
    education: str = ''
    achievements: str = ''
    target: str = ''
    tone: str = 'modern'
    template: str = 'modern'

@app.on_event('startup')
def startup():
    seed()

@app.get('/api/health')
def health():
    with connect() as conn:
        return {
            'status': 'ok',
            'database': str(DB_PATH),
            'resumes': conn.execute('SELECT COUNT(*) c FROM resumes').fetchone()['c'],
            'roles': conn.execute('SELECT COUNT(*) c FROM role_library').fetchone()['c']
        }

def normalize_resume_data(data: dict[str, Any]) -> dict[str, Any]:
    data = json.loads(json.dumps(data or {}))
    raw_skills = data.get('skills', [])
    if isinstance(raw_skills, dict):
        raw_skills = list(raw_skills.values())
    elif not isinstance(raw_skills, list):
        raw_skills = [raw_skills]
    skill_cells: list[str] = []
    for row in raw_skills:
        if isinstance(row, dict):
            items = row.get('items')
            if isinstance(items, list):
                cells = items
            elif isinstance(items, str):
                cells = re.split(r'[|,;•\n]', items)
            else:
                cells = [row.get('a', ''), row.get('b', ''), row.get('c', ''), row.get('name', ''), row.get('label', ''), row.get('value', '')]
        elif isinstance(row, str):
            cells = re.split(r'[|,;•\n]', row)
        else:
            cells = list(row) if isinstance(row, (list, tuple)) else []
        skill_cells.extend([str(x).strip() for x in cells if str(x).strip()])
    data['skills'] = [(skill_cells[i:i+3] + ['', '', ''])[:3] for i in range(0, len(skill_cells), 3)]
    for job in data.get('experience', []):
        if 'start_date' not in job or 'end_date' not in job:
            dates = job.get('dates', '')
            parts = re.split(r'\s+[–-]\s+', dates, maxsplit=1)
            job.setdefault('start_date', parts[0] if parts else '')
            job.setdefault('end_date', parts[1] if len(parts) > 1 else '')
    data.setdefault('custom_sections', [])
    return data


def row_resume(row):
    data = normalize_resume_data(loads(row['data_json']))
    return {'id': row['id'], 'name': row['name'], 'title': row['title'], 'template': row['template'], 'data': data, 'created_at': row['created_at'], 'updated_at': row['updated_at']}

@app.get('/api/resumes')
def resumes():
    with connect() as conn:
        return [row_resume(r) for r in conn.execute('SELECT * FROM resumes ORDER BY updated_at DESC')]

@app.get('/api/resumes/{resume_id}')
def resume(resume_id: int):
    with connect() as conn:
        r = conn.execute('SELECT * FROM resumes WHERE id=?', (resume_id,)).fetchone()
        if not r: raise HTTPException(404, 'Resume not found')
        return row_resume(r)

@app.post('/api/resumes')
def save_resume(payload: ResumeIn):
    stamp = now()
    with connect() as conn:
        cur = conn.execute('INSERT INTO resumes(name,title,template,data_json,created_at,updated_at) VALUES(?,?,?,?,?,?)',
                           (payload.name, payload.title, payload.template, dumps(normalize_resume_data(payload.data)), stamp, stamp))
        resume_id = cur.lastrowid
        if resume_id is None:
            raise HTTPException(500, 'Resume insert failed')
        conn.commit()
    return resume(int(resume_id))

@app.put('/api/resumes/{resume_id}')
def update_resume(resume_id: int, payload: ResumeIn):
    with connect() as conn:
        conn.execute('UPDATE resumes SET name=?, title=?, template=?, data_json=?, updated_at=? WHERE id=?',
                     (payload.name, payload.title, payload.template, dumps(normalize_resume_data(payload.data)), now(), resume_id))
    return resume(resume_id)

@app.delete('/api/resumes/{resume_id}')
def delete_resume(resume_id: int):
    with connect() as conn:
        conn.execute('DELETE FROM resumes WHERE id=?', (resume_id,))
    return {'deleted': resume_id}

@app.post('/api/sample-resume')
def load_sample_resume():
    return resume(insert_sample_resume())

@app.post('/api/resumes/{resume_id}/duplicate')
def duplicate(resume_id: int):
    r = resume(resume_id)
    payload = ResumeIn(name=r['name'] + ' - Copy', title=r['title'], template=r['template'], data=r['data'])
    return save_resume(payload)

@app.get('/api/roles')
def roles():
    with connect() as conn:
        rows = conn.execute('SELECT * FROM role_library ORDER BY category,title').fetchall()
        out = {}
        for r in rows:
            out.setdefault(r['category'], []).append({'title': r['title'], 'below_target': bool(r['below_target'])})
        return out

def row_document(row):
    return dict(row) | {'data': loads(row['data_json'])}

@app.get('/api/documents')
def documents(kind: Optional[str] = None):
    sql = 'SELECT * FROM documents' + (' WHERE kind=?' if kind else '') + ' ORDER BY updated_at DESC'
    args = (kind,) if kind else ()
    with connect() as conn:
        return [row_document(r) for r in conn.execute(sql, args)]

@app.get('/api/documents/{doc_id}')
def document(doc_id: int):
    with connect() as conn:
        r = conn.execute('SELECT * FROM documents WHERE id=?', (doc_id,)).fetchone()
        if not r: raise HTTPException(404, 'Document not found')
        return row_document(r)

@app.post('/api/documents')
def save_doc(payload: DocIn):
    stamp = now()
    with connect() as conn:
        cur = conn.execute('INSERT INTO documents(kind,title,resume_id,data_json,created_at,updated_at) VALUES(?,?,?,?,?,?)',
                           (payload.kind, payload.title, payload.resume_id, dumps(payload.data), stamp, stamp))
        doc_id = cur.lastrowid
        if doc_id is None: raise HTTPException(500, 'Document insert failed')
        conn.commit()
    return document(int(doc_id))

@app.put('/api/documents/{doc_id}')
def update_doc(doc_id: int, payload: DocIn):
    with connect() as conn:
        conn.execute('UPDATE documents SET kind=?, title=?, resume_id=?, data_json=?, updated_at=? WHERE id=?',
                     (payload.kind, payload.title, payload.resume_id, dumps(payload.data), now(), doc_id))
    return document(doc_id)

@app.delete('/api/documents/{doc_id}')
def delete_doc(doc_id: int):
    with connect() as conn:
        conn.execute('DELETE FROM documents WHERE id=?', (doc_id,))
    return {'deleted': doc_id}

def extract_text_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())

def extract_text_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    return '\n'.join(page.extract_text() or '' for page in reader.pages)

SECTION_ALIASES = {
    'summary': ['summary', 'professional summary', 'profile', 'objective'],
    'skills': ['areas of expertise', 'skills', 'core competencies', 'expertise'],
    'technical': ['technical proficiencies', 'technical skills', 'technologies', 'tools'],
    'experience': ['career experience', 'professional experience', 'experience', 'employment history', 'work history'],
    'education': ['education'],
    'certifications': ['certifications', 'certificates', 'licenses'],
    'additional': ['additional experience', 'additional information', 'projects']
}

def section_key(line: str) -> Optional[str]:
    clean = re.sub(r'[^a-zA-Z ]', '', line).strip().lower()
    for key, names in SECTION_ALIASES.items():
        if clean in names:
            return key
    return None

def split_sections(lines: list[str]) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {'header': []}
    current = 'header'
    for line in lines:
        key = section_key(line)
        if key:
            current = key
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(line)
    return sections

def bullet_lines(lines: list[str]) -> list[str]:
    out = []
    for line in lines:
        clean = re.sub(r'^[•\-*\u2022\s]+', '', line).strip()
        if clean:
            out.append(clean)
    return out

def parse_resume_text(text: str) -> dict[str, Any]:
    data = {'contact': {'name': '', 'title': '', 'email': '', 'phone': '', 'linkedin': '', 'portfolio': '', 'location': ''}, 'summary': '', 'skills': [], 'technical': {}, 'experience': [], 'education': [], 'certifications': [], 'additional': [], 'custom_sections': []}
    lines = [re.sub(r'\s+', ' ', l).strip() for l in text.splitlines() if l.strip()]
    sections = split_sections(lines)
    header = sections.get('header', [])[:8]
    email = re.search(r'[\w.+-]+@[\w.-]+', text)
    phone = re.search(r'(?:\+?1[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}', text)
    linkedin = re.search(r'https?://(?:www\.)?linkedin\.com/\S+|linkedin\.com/\S+', text, re.I)
    urls = re.findall(r'https?://\S+', text)
    if header:
        data['contact']['name'] = header[0][:100]
    if len(header) > 1 and not section_key(header[1]):
        data['contact']['title'] = header[1][:140]
    if email: data['contact']['email'] = email.group(0)
    if phone: data['contact']['phone'] = phone.group(0)
    if linkedin: data['contact']['linkedin'] = linkedin.group(0)
    for u in urls:
        if 'linkedin' not in u.lower():
            data['contact']['portfolio'] = u; break
    if len(header) > 2:
        maybe_location = [h for h in header if re.search(r'\b[A-Z]{2}\b|,', h) and '@' not in h and not re.search(r'\d{3}', h)]
        if maybe_location: data['contact']['location'] = maybe_location[-1][:100]
    summary_lines = sections.get('summary') or [l for l in header[2:] if '@' not in l and not re.search(r'\d{3}', l)]
    data['summary'] = ' '.join(summary_lines).strip()[:1200]
    skills = bullet_lines(sections.get('skills', []))
    if skills:
        cells = []
        for line in skills:
            cells += [x.strip() for x in re.split(r'[,|;•]', line) if x.strip()]
        data['skills'] = [(cells[i:i+3] + ['', '', ''])[:3] for i in range(0, len(cells), 3)]
    tech_lines = sections.get('technical', [])
    for line in tech_lines:
        if ':' in line:
            k, v = line.split(':', 1); data['technical'][k.strip()] = v.strip()
        else:
            data['technical'].setdefault('Tools', '')
            data['technical']['Tools'] = (data['technical']['Tools'] + ', ' + line).strip(', ')
    exp_lines = sections.get('experience', [])
    current_job = None
    for line in exp_lines:
        is_bullet = re.match(r'^[•\-*\u2022]', line) or len(line) > 95
        date_match = re.search(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec|January|February|March|April|June|July|August|September|October|November|December)?\s*\d{4})\s*[–-]\s*((?:Present|Current|\w+\s+)?\d{4}|Present|Current)', line, re.I)
        if not is_bullet and (date_match or len(line.split()) <= 12):
            if current_job: data['experience'].append(current_job)
            current_job = {'title': line, 'company': '', 'location': '', 'dates': date_match.group(0) if date_match else '', 'bullets': []}
        elif current_job:
            current_job['bullets'].append(re.sub(r'^[•\-*\u2022\s]+', '', line))
    if current_job: data['experience'].append(current_job)
    data['education'] = [{'degree': line, 'school': '', 'details': ''} for line in sections.get('education', [])]
    data['certifications'] = bullet_lines(sections.get('certifications', []))
    data['additional'] = bullet_lines(sections.get('additional', []))
    data['source_text'] = text[:20000]
    return normalize_resume_data(data)

def naive_parse_resume(text: str) -> dict[str, Any]:
    parsed = parse_resume_text(text)
    if not parsed['summary'] and not parsed['skills'] and not parsed['experience']:
        fallback = json.loads(json.dumps(DEFAULT_RESUME))
        fallback['source_text'] = text[:20000]
        return normalize_resume_data(fallback)
    return parsed

@app.post('/api/upload')
async def upload(file: UploadFile = File(...)):
    content = await file.read()
    suffix = Path(file.filename or '').suffix.lower()
    if suffix == '.pdf': text = extract_text_pdf(content)
    elif suffix in {'.docx', '.doc'}: text = extract_text_docx(content)
    else: raise HTTPException(400, 'Upload a PDF or Word document')
    path = UPLOAD_DIR / (file.filename or 'resume-upload')
    path.write_bytes(content)
    parsed = naive_parse_resume(text)
    return {'filename': file.filename or 'resume-upload', 'original_text': text[:30000], 'parsed': parsed}

@app.post('/api/import-confirm')
def import_confirm(payload: ImportConfirmIn):
    data = normalize_resume_data(payload.data)
    name = data.get('contact', {}).get('name') or Path(payload.filename).stem or 'Imported Resume'
    title = data.get('contact', {}).get('title') or 'Imported Resume'
    return save_resume(ResumeIn(name=f'Imported - {name}', title=title, template=payload.template, data=data))

def keywords(text: str) -> set[str]:
    words = re.findall(r'\b[A-Za-z][A-Za-z0-9+#.-]{2,}\b', text.lower())
    stop = {'the','and','for','with','you','are','will','that','this','from','our','your','have','has','into','using','work','team','role','job'}
    return {w for w in words if w not in stop}

def claude(prompt: str) -> Optional[str]:
    key = os.getenv('ANTHROPIC_API_KEY') or os.getenv('ANTHROPIC_TOKEN')
    if not key: return None
    try:
        resp = requests.post('https://api.anthropic.com/v1/messages', headers={'x-api-key': key, 'anthropic-version': '2023-06-01', 'content-type': 'application/json'}, json={'model': os.getenv('ANTHROPIC_MODEL', 'claude-sonnet-4-20250514'), 'max_tokens': 1800, 'messages': [{'role': 'user', 'content': prompt}]}, timeout=60)
        resp.raise_for_status()
        return resp.json()['content'][0]['text']
    except Exception:
        return None

@app.post('/api/tailor')
def tailor(payload: TailorIn):
    r = resume(payload.resume_id)
    resume_text = json.dumps(r['data'])
    missing = sorted(list(keywords(payload.job_description) - keywords(resume_text)))[:30]
    score = max(0, min(100, round(100 * (1 - len(missing) / max(30, len(keywords(payload.job_description)))))))
    generated = claude('Analyze this resume JSON against this job description. Return concise tailored summary, missing keyword recommendations, and revised bullets.\nRESUME:\n' + resume_text + '\nJOB:\n' + payload.job_description)
    return {'ats_score': score, 'missing_keywords': missing, 'suggestions': generated or 'Claude API key not configured. Keyword analysis completed locally. Add the missing keywords where truthful and relevant.', 'tailored_resume': r['data']}

def contact_from_resume(resume_id: Optional[int]) -> dict[str, str]:
    if not resume_id:
        return {'name': '', 'contact': ''}
    try:
        c = resume(resume_id)['data'].get('contact', {})
        contact = ' | '.join([x for x in [c.get('email'), c.get('phone'), c.get('location'), c.get('linkedin')] if x])
        return {'name': c.get('name', ''), 'contact': contact}
    except Exception:
        return {'name': '', 'contact': ''}

def resume_context(resume_id: Optional[int]) -> dict[str, Any]:
    if not resume_id:
        return {'contact': {}, 'summary': '', 'skills': [], 'achievements': [], 'roles': []}
    try:
        data = resume(resume_id)['data']
    except Exception:
        return {'contact': {}, 'summary': '', 'skills': [], 'achievements': [], 'roles': []}
    achievements: list[str] = []
    roles: list[str] = []
    for job in data.get('experience', [])[:5]:
        title = ' / '.join([x for x in [job.get('title'), job.get('company')] if x])
        if title:
            roles.append(title)
        achievements.extend([str(b) for b in job.get('bullets', []) if str(b).strip()][:3])
    skills = [cell for row in normalize_resume_data(data).get('skills', []) for cell in row if cell]
    return {
        'contact': data.get('contact', {}),
        'summary': data.get('summary', ''),
        'skills': skills[:36],
        'achievements': achievements[:18],
        'roles': roles,
        'technical': data.get('technical', {}),
        'education': data.get('education', []),
        'certifications': data.get('certifications', []),
    }

def professional_letter_prompt(kind: str, details: dict[str, Any], sender_name: str, sender_contact: str, context: dict[str, Any]) -> str:
    level_hint = 'senior/executive-level, polished, confident, commercially mature, and appropriate for roles at or above $115k when the job description implies senior scope'
    shared_rules = """Quality bar:
- Produce a final, ready-to-send business letter, not advice, notes, markdown, bullets, or a template.
- Use proper US business letter format: sender block, date, recipient/company line when available, salutation, body, professional closing, sender name.
- Write 3-4 strong paragraphs after the salutation. Each paragraph must have a clear job-specific purpose.
- Use company, role, job-description language, interview details, and linked-resume evidence when provided.
- Pull specific achievements, skills, certifications, technologies, leadership scope, operating environments, and measurable outcomes from the linked resume context. Do not invent numbers; if metrics are missing, describe impact credibly without fake statistics.
- Avoid generic phrases such as "dynamic professional," "proven track record," "I am excited to apply," and "perfect fit." Sound like a senior professional who has actually done the work.
- If information is missing, write around it naturally instead of leaving bracket placeholders, except for a truly unavailable recipient address line.
- Return only the finished letter text."""
    if kind == 'cover_letter':
        task = """Specific cover-letter structure:
1. Opening: name the company and target role if supplied, state a crisp value proposition tied to the job description, and show informed interest in the organization or problem space.
2. Body paragraph 1: map the most relevant resume achievements and skills to the employer's stated needs.
3. Body paragraph 2: reinforce seniority with leadership, operational judgment, stakeholder management, security/risk, scale, business outcomes, or technical depth as applicable.
4. Closing: confident request for a conversation and concise statement of the value the candidate would bring."""
    else:
        task = """Specific thank-you-letter structure:
1. Opening: thank the interviewer by name when supplied and reference the company, role, and interview date/context.
2. Body paragraph 1: mention the provided talking points and connect them to relevant resume achievements or skills.
3. Body paragraph 2: reinforce fit, senior judgment, and near-term business value based on what was discussed.
4. Closing: reaffirm interest, appreciation, and readiness for next steps without sounding needy or generic."""
    return f"""You are an elite executive resume writer and career strategist.

Tone: {level_hint}. Be specific, concise, polished, and credible.

{shared_rules}

{task}

Sender name: {sender_name or '[Your Name]'}
Sender contact: {sender_contact or '[Email] | [Phone] | [Location]'}
Details JSON:
{json.dumps(details, indent=2)}
Linked resume context JSON:
{json.dumps(context, indent=2)}
"""

def today_long() -> str:
    return datetime.utcnow().strftime('%B %-d, %Y') if os.name != 'nt' else datetime.utcnow().strftime('%B %#d, %Y')

def cover_template(company='[Company]', manager='Hiring Manager', sender_name='', sender_contact=''):
    name = sender_name or '[Your Name]'
    contact = sender_contact or '[Email] | [Phone] | [Location]'
    return f"{name}\n{contact}\n\n{today_long()}\n\n{company}\n[Company Address]\n\nDear {manager or 'Hiring Manager'},\n\nI am writing to express my interest in the [Position Title] role with {company}. My background in [relevant discipline], combined with hands-on experience in [key skill area], aligns well with the needs of your team.\n\nIn previous roles, I have delivered measurable results by [achievement one], [achievement two], and [achievement three]. I bring a practical, professional approach to solving problems, supporting stakeholders, and improving operational outcomes.\n\nI would welcome the opportunity to discuss how my experience can support {company}'s goals. Thank you for your time and consideration.\n\nSincerely,\n{name}"

def thank_template(company='[Company]', interviewer='Hiring Manager', position='[Position Title]', interview_date='[Interview Date]', sender_name='', sender_contact=''):
    name = sender_name or '[Your Name]'
    contact = sender_contact or '[Email] | [Phone] | [Location]'
    return f"{name}\n{contact}\n\n{today_long()}\n\nDear {interviewer or 'Hiring Manager'},\n\nThank you for taking the time to speak with me on {interview_date} about the {position} opportunity at {company}. I appreciated learning more about the role, the team, and the priorities ahead.\n\nOur conversation reinforced my interest in the position, particularly the discussion around [key topic]. My background in [relevant experience] would allow me to contribute quickly and thoughtfully.\n\nThank you again for your time and consideration. I look forward to the possibility of continuing the conversation.\n\nSincerely,\n{name}"

@app.post('/api/cover-letter')
def cover(payload: CoverIn):
    linked = contact_from_resume(payload.resume_id)
    sender_name = payload.sender_name or linked['name']
    sender_contact = payload.sender_contact or linked['contact']
    details = payload.model_dump()
    prompt = professional_letter_prompt('cover_letter', details, sender_name, sender_contact, resume_context(payload.resume_id))
    text = claude(prompt) or cover_template(payload.company, payload.hiring_manager, sender_name, sender_contact)
    return save_doc(DocIn(kind='cover_letter', title=f'Cover Letter - {payload.company or "Draft"}', resume_id=payload.resume_id, data={'content': text, 'sender_name': sender_name, 'sender_contact': sender_contact, **payload.model_dump()}))

@app.post('/api/thank-you')
def thank_you(payload: ThankYouIn):
    linked = contact_from_resume(payload.resume_id)
    sender_name = payload.sender_name or linked['name']
    sender_contact = payload.sender_contact or linked['contact']
    details = payload.model_dump()
    prompt = professional_letter_prompt('thank_you', details, sender_name, sender_contact, resume_context(payload.resume_id))
    text = claude(prompt) or thank_template(payload.company, payload.interviewer, payload.position, payload.interview_date, sender_name, sender_contact)
    return save_doc(DocIn(kind='thank_you', title=f'Thank You - {payload.company or "Draft"}', resume_id=payload.resume_id, data={'content': text, 'sender_name': sender_name, 'sender_contact': sender_contact, **payload.model_dump()}))

ROLE_SPECIALIZATIONS: dict[str, dict[str, Any]] = {
    'azure administrator': {
        'skills': ['Azure Administration', 'ARM Templates', 'Virtual Networks', 'Network Security Groups', 'Entra ID', 'RBAC', 'Azure Monitor', 'Cost Management', 'Azure Policy', 'Backup / Recovery', 'PowerShell / CLI', 'AZ-104 Readiness'],
        'technical': {'Azure Infrastructure': 'ARM templates, VNets, subnets, NSGs, route tables, Azure Bastion, storage accounts', 'Identity': 'Microsoft Entra ID, RBAC, conditional access, privileged access reviews', 'Operations': 'Azure Monitor, Log Analytics, alerts, backup, cost management, Azure Policy'},
        'certs': ['Microsoft Certified: Azure Administrator Associate (AZ-104)', 'Microsoft Azure Fundamentals'],
        'verbs': ['administered Azure landing-zone resources', 'standardized RBAC and policy assignments', 'improved visibility with Azure Monitor and Log Analytics']
    },
    'network administrator': {
        'skills': ['Cisco IOS', 'VLANs / Trunking', 'BGP / OSPF', 'pfSense Firewalls', 'Wireshark Analysis', 'VPN Administration', 'Network Segmentation', 'Switching / Routing', 'DHCP / DNS', 'Wireless LAN', 'CCNA Readiness', 'Documentation'],
        'technical': {'Network Platforms': 'Cisco IOS, pfSense, UniFi, managed switching, wireless controllers', 'Protocols': 'VLANs, STP, BGP, OSPF, IPsec, DHCP, DNS, NAT, ACLs', 'Tools': 'Wireshark, SNMP monitoring, syslog, NetFlow, configuration backup'},
        'certs': ['Cisco Certified Network Associate (CCNA)', 'CompTIA Network+'],
        'verbs': ['maintained routed and switched infrastructure', 'segmented network services with VLANs and firewall rules', 'resolved packet-level issues using Wireshark and monitoring data']
    },
    'endpoint engineer': {
        'skills': ['Microsoft Intune', 'Windows Autopilot', 'Compliance Policies', 'Configuration Profiles', 'MD-102 Readiness', 'SCCM Co-management', 'Endpoint Security', 'App Deployment', 'Patch Rings', 'Conditional Access', 'Device Enrollment', 'PowerShell'],
        'technical': {'Endpoint Management': 'Intune, Autopilot, compliance policies, configuration profiles, enrollment restrictions', 'Co-management': 'SCCM/MECM, collections, application packaging, update rings, device inventory', 'Security': 'Defender for Endpoint, BitLocker, baseline policies, conditional access signals'},
        'certs': ['Microsoft 365 Certified: Endpoint Administrator Associate (MD-102)', 'Microsoft Certified: Security, Compliance, and Identity Fundamentals'],
        'verbs': ['engineered Intune and Autopilot endpoint standards', 'deployed compliance and configuration profiles', 'coordinated SCCM co-management and application rollout']
    },
    'm365 engineer': {
        'skills': ['Exchange Online', 'SharePoint Online', 'Teams Administration', 'DLP Policies', 'MS-102 Readiness', 'Power Platform', 'Purview', 'Entra ID', 'Mailbox Migration', 'Retention Policies', 'Secure Collaboration', 'PowerShell'],
        'technical': {'Microsoft 365': 'Exchange Online, SharePoint Online, Teams admin center, OneDrive, Microsoft Purview', 'Security / Compliance': 'DLP policies, retention, eDiscovery, sensitivity labels, conditional access', 'Automation': 'PowerShell, Power Platform, reporting, service-health monitoring'},
        'certs': ['Microsoft 365 Certified: Administrator Expert (MS-102)', 'Microsoft 365 Fundamentals'],
        'verbs': ['administered Microsoft 365 collaboration workloads', 'implemented DLP and retention policies', 'automated Exchange Online and Teams administration tasks']
    },
}
ROLE_SPECIALIZATIONS['microsoft 365 engineer'] = ROLE_SPECIALIZATIONS['m365 engineer']

CATEGORY_PROFILES: dict[str, dict[str, Any]] = {
    'Network Administration': {'skills': ['Routing / Switching', 'Firewall Administration', 'VPN Operations', 'Network Monitoring', 'Incident Response', 'Documentation', 'Wireless Networks', 'DNS / DHCP', 'Change Control', 'Vendor Coordination', 'Capacity Planning', 'Security Segmentation'], 'certs': ['CompTIA Network+', 'CCNA'], 'technical': {'Network': 'routing, switching, VLANs, firewall policy, VPN, wireless, DNS/DHCP', 'Operations': 'monitoring, diagrams, change control, vendor escalation, incident response'}},
    'Systems Administration': {'skills': ['Windows Server', 'Active Directory', 'Group Policy', 'Virtualization', 'Backup / Recovery', 'Patch Management', 'PowerShell', 'Server Monitoring', 'Identity Administration', 'Storage', 'Documentation', 'Incident Resolution'], 'certs': ['Microsoft Certified: Windows Server Hybrid Administrator', 'CompTIA Server+'], 'technical': {'Systems': 'Windows Server, AD DS, Group Policy, Hyper-V/VMware, storage, backup platforms', 'Automation': 'PowerShell, scheduled maintenance, monitoring, patching, documentation'}},
    'Endpoint / Microsoft 365': {'skills': ['Microsoft Intune', 'Microsoft 365 Admin', 'Endpoint Compliance', 'Autopilot', 'Exchange Online', 'Teams Administration', 'SharePoint Online', 'PowerShell', 'Device Lifecycle', 'Security Baselines', 'User Support', 'Reporting'], 'certs': ['Microsoft 365 Certified: Endpoint Administrator Associate', 'Microsoft 365 Fundamentals'], 'technical': {'Modern Workplace': 'Intune, Microsoft 365 admin centers, Autopilot, Teams, Exchange Online, SharePoint', 'Endpoint': 'configuration profiles, compliance policy, app deployment, update rings, device inventory'}},
    'Cloud & Hybrid': {'skills': ['Azure Infrastructure', 'Hybrid Identity', 'Cloud Networking', 'RBAC', 'Monitoring', 'IaC Templates', 'Backup / Recovery', 'Cost Control', 'Policy Governance', 'Migration Support', 'Automation', 'Security Controls'], 'certs': ['Microsoft Certified: Azure Administrator Associate', 'Azure Fundamentals'], 'technical': {'Cloud': 'Azure compute, storage, networking, Entra ID, RBAC, Azure Monitor, backup', 'Hybrid': 'site connectivity, identity synchronization, migration planning, policy governance'}},
    'Security': {'skills': ['Security Monitoring', 'Incident Triage', 'Vulnerability Management', 'Endpoint Protection', 'SIEM Review', 'Access Controls', 'Policy Enforcement', 'Risk Documentation', 'Phishing Response', 'Log Analysis', 'Hardening', 'Compliance Support'], 'certs': ['CompTIA Security+', 'Microsoft Security, Compliance, and Identity Fundamentals'], 'technical': {'Security Operations': 'SIEM, EDR, vulnerability scanners, log review, incident documentation', 'Controls': 'MFA, access reviews, endpoint hardening, policy enforcement, remediation tracking'}},
    'Executive & Administrative Support': {'skills': ['Executive Calendar Management', 'Travel Coordination', 'Board Materials', 'Confidential Communication', 'Meeting Logistics', 'Expense Reporting', 'Stakeholder Liaison', 'Office Operations', 'Process Improvement', 'Vendor Coordination', 'Document Preparation', 'Priority Management'], 'certs': ['Certified Administrative Professional (CAP)', 'Microsoft Office Specialist'], 'technical': {'Business Tools': 'Microsoft Office, Outlook, Teams, Zoom, expense platforms, CRM/ERP coordination', 'Administration': 'calendar workflows, board packets, travel, correspondence, office procedures'}},
    'Accounting & Finance': {'skills': ['Month-End Close', 'Account Reconciliation', 'Financial Reporting', 'General Ledger', 'AP / AR', 'Payroll Coordination', 'Budget Support', 'Variance Analysis', 'Tax Documentation', 'Audit Support', 'Excel Modeling', 'Internal Controls'], 'certs': ['QuickBooks ProAdvisor', 'Microsoft Excel Expert'], 'technical': {'Finance Systems': 'QuickBooks, ERP/accounting platforms, Excel, payroll systems, bank portals', 'Accounting': 'GL, reconciliations, close, AP/AR, reporting, controls, audit support'}},
    'CPA Firm': {'skills': ['Tax Preparation', 'Client Accounting Services', 'Audit Coordination', 'Practice Workflow', 'Engagement Tracking', 'Client Communication', 'Document Management', 'Tax Operations', 'Deadline Management', 'Billing Support', 'Compliance Documentation', 'Workflow Improvement'], 'certs': ['IRS Annual Filing Season Program', 'QuickBooks ProAdvisor'], 'technical': {'Practice Tools': 'tax software, QuickBooks, document portals, workflow management, e-signature tools', 'Client Service': 'engagement setup, organizer tracking, deliverables, deadline control'}},
    'HR & Recruiting': {'skills': ['Candidate Coordination', 'ATS Administration', 'Onboarding', 'Benefits Support', 'Employee Relations Intake', 'Compliance Documentation', 'Interview Scheduling', 'Offer Processing', 'HRIS Data Quality', 'Policy Communication', 'Reporting', 'Confidential Records'], 'certs': ['SHRM-CP Coursework', 'HR Management Certificate'], 'technical': {'HR Systems': 'ATS, HRIS, benefits portals, background-check platforms, Microsoft Office', 'People Operations': 'onboarding, candidate workflows, employee records, compliance tracking'}},
    'Legal & Professional Services': {'skills': ['Matter Management', 'Legal Drafting', 'eFiling', 'Discovery Support', 'Calendar / Docket Control', 'Client Communication', 'Document Review', 'Case Preparation', 'Billing Support', 'Confidential Records', 'Research Support', 'Deadline Management'], 'certs': ['Paralegal Certificate', 'Notary Public'], 'technical': {'Legal Tools': 'case management systems, eFiling portals, document management, Microsoft Office', 'Practice Support': 'pleadings, discovery, docketing, client files, billing coordination'}},
    'Healthcare': {'skills': ['Practice Operations', 'Patient Scheduling', 'Insurance Verification', 'HIPAA Compliance', 'Medical Records', 'Front Office Leadership', 'Revenue Cycle Support', 'Provider Coordination', 'Patient Communication', 'EHR Workflows', 'Referral Tracking', 'Staff Coordination'], 'certs': ['Certified Medical Administrative Assistant', 'HIPAA Training'], 'technical': {'Healthcare Systems': 'EHR/EMR, scheduling, insurance portals, referral platforms, patient communication tools', 'Operations': 'front desk, records, revenue cycle support, provider schedules, HIPAA controls'}},
    'Real Estate': {'skills': ['Lease Administration', 'Tenant Relations', 'Property Operations', 'Vendor Coordination', 'Transaction Support', 'Closing Documentation', 'Budget Tracking', 'Maintenance Coordination', 'Compliance Records', 'CRM Updates', 'Market Documentation', 'Client Communication'], 'certs': ['Real Estate License Coursework', 'Property Management Certificate'], 'technical': {'Real Estate Tools': 'property management systems, MLS/CRM tools, e-signature platforms, Excel', 'Operations': 'leases, closings, vendor work orders, tenant files, transaction checklists'}},
    'Operations & Business Support': {'skills': ['Project Coordination', 'Process Improvement', 'Vendor Management', 'Compliance Tracking', 'Client Service', 'Reporting', 'SOP Documentation', 'Cross-Functional Support', 'Scheduling', 'Budget Support', 'Risk Follow-Up', 'Service Delivery'], 'certs': ['Certified Associate in Project Management (CAPM)', 'Lean Six Sigma Yellow Belt'], 'technical': {'Operations Tools': 'Microsoft 365, project trackers, CRM, ticketing/workflow tools, reporting dashboards', 'Business Support': 'SOPs, vendor follow-up, compliance logs, client communication, process metrics'}},
}

ROLE_FOCUS_LIBRARY: dict[str, list[str]] = {
    'network': ['Cisco IOS', 'VLAN design', 'BGP/OSPF routing', 'pfSense policy', 'Wireshark packet analysis', 'CCNA-level troubleshooting'],
    'systems': ['Windows Server', 'Active Directory', 'Group Policy', 'PowerShell automation', 'backup validation', 'patch governance'],
    'endpoint': ['Microsoft Intune', 'Windows Autopilot', 'compliance policies', 'configuration profiles', 'MD-102 readiness', 'SCCM co-management'],
    'm365': ['Exchange Online', 'SharePoint Online', 'Teams administration', 'DLP policies', 'MS-102 readiness', 'Power Platform'],
    'azure': ['ARM templates', 'VNets', 'NSGs', 'Entra ID', 'AZ-104 readiness', 'Azure RBAC'],
    'security': ['SIEM triage', 'EDR response', 'vulnerability remediation', 'access reviews', 'hardening baselines', 'incident documentation'],
    'executive': ['executive calendar control', 'board packet preparation', 'travel logistics', 'confidential correspondence', 'meeting orchestration', 'expense reconciliation'],
    'accounting': ['general ledger', 'month-end close', 'account reconciliations', 'financial reporting', 'audit support', 'Excel analysis'],
    'tax': ['1040/1120/1065 returns', 'tax planning', 'workpaper review', 'client advisory', 'tax research', 'deadline management'],
    'bookkeeping': ['QuickBooks', 'bank reconciliations', 'AP/AR processing', 'payroll coordination', 'sales tax filings', 'monthly financial packages'],
    'payroll': ['payroll processing', 'timecard audits', 'benefit deductions', 'tax deposits', 'year-end W-2s', 'payroll compliance'],
    'audit': ['audit workpapers', 'PBC tracking', 'substantive testing', 'control walkthroughs', 'sampling documentation', 'engagement binders'],
    'hr': ['HRIS data quality', 'onboarding workflows', 'benefits administration', 'employee records', 'policy communication', 'compliance tracking'],
    'recruiting': ['ATS coordination', 'interview scheduling', 'candidate pipeline updates', 'offer processing', 'background checks', 'hiring manager follow-up'],
    'legal': ['matter management', 'eFiling', 'discovery support', 'pleading preparation', 'docket control', 'client file maintenance'],
    'healthcare': ['EHR workflows', 'patient scheduling', 'insurance verification', 'HIPAA controls', 'referral tracking', 'revenue-cycle support'],
    'real_estate': ['lease administration', 'tenant relations', 'property budgets', 'vendor work orders', 'transaction files', 'closing coordination'],
    'operations': ['process improvement', 'SOP documentation', 'vendor management', 'project tracking', 'compliance logs', 'service reporting'],
    'finance': ['variance analysis', 'cash forecasting', 'budget modeling', 'management reporting', 'controls review', 'board packages'],
}

TITLE_FOCUS_RULES: list[tuple[str, str]] = [
    ('azure', 'azure'), ('cloud', 'azure'), ('hybrid', 'azure'), ('identity', 'azure'), ('entra', 'azure'),
    ('network', 'network'), ('unified communications', 'network'),
    ('endpoint', 'endpoint'), ('intune', 'endpoint'), ('modern workplace', 'endpoint'), ('euc', 'endpoint'),
    ('microsoft 365', 'm365'), ('m365', 'm365'),
    ('security', 'security'), ('soc', 'security'),
    ('systems', 'systems'), ('server', 'systems'), ('windows', 'systems'), ('infrastructure', 'systems'),
    ('tax', 'tax'), ('bookkeeper', 'bookkeeping'), ('payroll', 'payroll'), ('audit', 'audit'), ('accountant', 'accounting'), ('controller', 'finance'), ('financial', 'finance'), ('finance', 'finance'),
    ('recruiting', 'recruiting'), ('talent acquisition', 'recruiting'), ('hr ', 'hr'), ('benefits', 'hr'), ('onboarding', 'hr'), ('employee relations', 'hr'),
    ('legal', 'legal'), ('paralegal', 'legal'), ('litigation', 'legal'),
    ('medical', 'healthcare'), ('clinical', 'healthcare'), ('patient', 'healthcare'),
    ('property', 'real_estate'), ('real estate', 'real_estate'), ('transaction', 'real_estate'), ('closing', 'real_estate'),
    ('executive', 'executive'), ('board', 'executive'), ('chief of staff', 'executive'), ('office manager', 'executive'),
]

CATEGORY_FOCUS = {
    'Network Administration': 'network', 'Systems Administration': 'systems', 'Endpoint / Microsoft 365': 'm365', 'Cloud & Hybrid': 'azure', 'Security': 'security',
    'Executive & Administrative Support': 'executive', 'Accounting & Finance': 'accounting', 'CPA Firm': 'tax', 'HR & Recruiting': 'hr', 'Legal & Professional Services': 'legal',
    'Healthcare': 'healthcare', 'Real Estate': 'real_estate', 'Operations & Business Support': 'operations'
}

ROLE_OVERRIDES = {
    'tax manager': ['Tax Planning', 'Review Workpapers', 'Client Advisory', 'Entity Returns', 'Tax Research', 'Staff Coaching'],
    'controller': ['Close Leadership', 'Internal Controls', 'Board Reporting', 'Cash Forecasting', 'Audit Leadership', 'Team Management'],
    'paralegal': ['Discovery', 'Pleadings', 'Legal Research', 'Trial Preparation', 'eFiling', 'Matter Management'],
    'property manager': ['Tenant Relations', 'Lease Compliance', 'Maintenance Coordination', 'Property Budgets', 'Vendor Oversight', 'Occupancy Reporting'],
    'chief of staff': ['Executive Rhythm', 'Strategic Initiatives', 'Board Readiness', 'Operating Cadence', 'Decision Support', 'Cross-Functional Alignment'],
}


def role_category(title: str, requested_category: Optional[str] = None) -> str:
    if requested_category in ROLE_LIBRARY:
        return requested_category
    for category, titles in ROLE_LIBRARY.items():
        if title in titles:
            return category
    lowered = title.lower()
    for category, titles in ROLE_LIBRARY.items():
        if any(t.lower() == lowered for t in titles):
            return category
    return 'Operations & Business Support'


def focus_key_for_role(title: str, category: str) -> str:
    lowered = f' {title.lower()} '
    for needle, key in TITLE_FOCUS_RULES:
        if needle in lowered:
            return key
    return CATEGORY_FOCUS.get(category, 'operations')


def role_profile(title: str, requested_category: Optional[str] = None) -> dict[str, Any]:
    lowered = title.lower().strip()
    if lowered in ROLE_SPECIALIZATIONS:
        return json.loads(json.dumps(ROLE_SPECIALIZATIONS[lowered]))
    category = role_category(title, requested_category)
    profile = json.loads(json.dumps(CATEGORY_PROFILES.get(category, CATEGORY_PROFILES['Operations & Business Support'])))
    focus_key = focus_key_for_role(title, category)
    role_terms = ROLE_FOCUS_LIBRARY[focus_key]
    for key, extra in ROLE_OVERRIDES.items():
        if key in lowered:
            role_terms = extra + [x for x in role_terms if x not in extra]
            break
    title_tokens = [w for w in re.split(r'[^A-Za-z0-9]+', title) if len(w) > 2 and w.lower() not in {'senior', 'administrator', 'engineer', 'manager', 'coordinator', 'assistant', 'specialist'}]
    distinctive = [f'{token} Operations' for token in title_tokens[:2]] + [f'{title} Stakeholder Support']
    profile['skills'] = (role_terms + distinctive + [x for x in profile['skills'] if x not in role_terms and x not in distinctive])[:15]
    profile['technical']['Role Focus'] = f'{title}: {", ".join(role_terms[:6])}'
    profile['technical']['Execution Context'] = f'{category} standards, role-specific reporting, documentation, stakeholder communication, and audit-ready handoffs'
    profile['certs'] = list(dict.fromkeys(profile.get('certs', []) + role_certifications(title, focus_key)))[:4]
    profile['verbs'] = role_verbs(title, role_terms, category)
    return profile


def role_certifications(title: str, focus_key: str) -> list[str]:
    certs = {
        'network': ['Cisco Certified Network Associate (CCNA)', 'CompTIA Network+'],
        'systems': ['Microsoft Certified: Windows Server Hybrid Administrator Associate', 'CompTIA Server+'],
        'endpoint': ['Microsoft 365 Certified: Endpoint Administrator Associate (MD-102)'],
        'm365': ['Microsoft 365 Certified: Administrator Expert (MS-102)', 'Microsoft 365 Fundamentals'],
        'azure': ['Microsoft Certified: Azure Administrator Associate (AZ-104)', 'Microsoft Azure Fundamentals'],
        'security': ['CompTIA Security+', 'Microsoft Security, Compliance, and Identity Fundamentals'],
        'executive': ['Certified Administrative Professional (CAP)'],
        'accounting': ['QuickBooks ProAdvisor', 'Microsoft Excel Expert'],
        'tax': ['IRS Annual Filing Season Program', 'QuickBooks ProAdvisor'],
        'bookkeeping': ['QuickBooks ProAdvisor'],
        'payroll': ['Fundamental Payroll Certification (FPC)'],
        'audit': ['Audit Analytics / Workpaper Training'],
        'hr': ['SHRM-CP Coursework'],
        'recruiting': ['LinkedIn Recruiter / ATS Workflow Training'],
        'legal': ['Paralegal Certificate'],
        'healthcare': ['HIPAA Training', 'Certified Medical Administrative Assistant'],
        'real_estate': ['Property Management Certificate', 'Real Estate License Coursework'],
        'operations': ['CAPM', 'Lean Six Sigma Yellow Belt'],
        'finance': ['Microsoft Excel Expert', 'FP&A Modeling Coursework'],
    }
    return certs.get(focus_key, [])


def role_verbs(title: str, terms: list[str], category: str) -> list[str]:
    return [
        f'owned {title.lower()} delivery across {terms[0]}, {terms[1]}, and {terms[2]}',
        f'improved {category.lower()} outcomes by standardizing {terms[3]}, {terms[4]}, and recurring status reporting',
        f'converted {title.lower()} requirements into documented workflows, stakeholder-ready updates, and measurable service improvements',
    ]


def fallback_role_resume(payload: RoleBuildIn) -> dict[str, Any]:
    category = role_category(payload.title, payload.category)
    profile = role_profile(payload.title, category)
    skills = profile['skills'][:15]
    verbs = profile.get('verbs') or role_verbs(payload.title, skills, category)
    return normalize_resume_data({
        'contact': {'name': '', 'title': payload.title, 'email': '', 'phone': '', 'linkedin': '', 'portfolio': '', 'location': ''},
        'summary': f'{payload.title} with focused experience in {category.lower()} environments, combining {skills[0]}, {skills[1]}, and {skills[2]} with disciplined documentation, stakeholder partnership, and measurable execution. Known for applying {skills[3]}, {skills[4]}, and {skills[5]} to improve reliability, control quality, and operational follow-through for {payload.industry} organizations.',
        'skills': [(skills[i:i+3] + ['', '', ''])[:3] for i in range(0, len(skills), 3)],
        'technical': profile['technical'],
        'experience': [
            {'title': payload.title, 'company': 'Confidential Organization', 'location': '', 'dates': '2022 – Present', 'bullets': [
                f'{verbs[0].capitalize()} while maintaining accurate documentation, status reporting, and escalation discipline.',
                f'{verbs[1].capitalize()} using role-appropriate tools, controls, and communication routines.',
                f'Partnered with leadership, peers, vendors, and end users to prioritize work, remove blockers, and improve {payload.title.lower()} outcomes.'
            ]},
            {'title': f'{payload.title} / Operations Specialist', 'company': 'Professional Services Organization', 'location': '', 'dates': '2019 – 2022', 'bullets': [
                f'Supported {payload.title.lower()} functions across {skills[0].lower()}, {skills[1].lower()}, issue tracking, documentation, and stakeholder follow-up.',
                f'Built checklists, templates, and reporting cadences around {skills[2].lower()} and {skills[3].lower()} that improved consistency and reduced preventable rework.',
                f'Handled confidential information, deadlines, and competing priorities with practical judgment and service-focused execution.'
            ]}
        ],
        'education': [{'degree': f'Professional development aligned to {payload.title}', 'school': '', 'details': payload.industry}],
        'certifications': profile.get('certs', []),
        'additional': [f'Selected projects include {skills[0].lower()}, {skills[1].lower()}, {skills[2].lower()}, workflow cleanup, reporting improvement, and stakeholder-ready documentation.'],
        'custom_sections': [{'title': 'Role-Specific Impact', 'bullets': [f'Applied {skills[0]}, {skills[1]}, and {skills[2]} to improve execution quality.', f'Maintained clean handoffs and audit-ready documentation for {payload.title.lower()} responsibilities in {category.lower()} environments.']}]
    })

@app.post('/api/role-build')
def role_build(payload: RoleBuildIn):
    prompt = '''Create premium resume content JSON for the requested target role. Return ONLY JSON matching this schema:
{"contact":{"name":"","title":"","email":"","phone":"","linkedin":"","portfolio":"","location":""},"summary":"","skills":[["","",""]],"technical":{"Category":"items"},"experience":[{"title":"","company":"","location":"","dates":"","bullets":[""]}],"education":[{"degree":"","school":"","details":""}],"certifications":[],"additional":[],"custom_sections":[{"title":"","bullets":[""]}]}
Requirements: 12-18 role-specific Areas of Expertise in three-column rows, strong executive summary, credible achievement-oriented bullets, and role-specific technical/tool categories. Keep contact fields blank except title. The content must be unmistakably specific to the exact requested title, not broad generic IT/operations language. For example: Azure Administrator must include ARM templates, VNets, NSGs, Entra ID, AZ-104, Azure Monitor, RBAC, and Cost Management; Network Administrator must include Cisco IOS, VLANs, BGP/OSPF, pfSense, Wireshark, CCNA, and network segmentation; Endpoint Engineer must include Intune, Autopilot, compliance policies, configuration profiles, MD-102, and SCCM co-management; M365 Engineer must include Exchange Online, SharePoint, Teams admin, DLP policies, MS-102, and Power Platform. Non-IT roles must avoid IT-only technical sections unless the requested title is technical.'''
    prompt += '\nROLE PROFILE JSON:' + json.dumps(role_profile(payload.title, payload.category), indent=2)
    prompt += '\nREQUEST:' + payload.model_dump_json()
    generated = claude(prompt)
    parsed = extract_json_object(generated or '')
    if parsed:
        parsed.setdefault('contact', {})
        parsed['contact'] = {**{'name': '', 'title': '', 'email': '', 'phone': '', 'linkedin': '', 'portfolio': '', 'location': ''}, **parsed.get('contact', {})}
        parsed['contact']['title'] = payload.title
        return {'resume': normalize_resume_data(parsed), 'raw_generation': generated}
    return {'resume': fallback_role_resume(payload), 'raw_generation': generated}

def extract_json_object(text: str) -> Optional[dict[str, Any]]:
    if not text:
        return None
    start = text.find('{'); end = text.rfind('}')
    if start == -1 or end == -1 or end <= start:
        return None
    try:
        return json.loads(text[start:end+1])
    except Exception:
        return None

def fallback_personal_resume(payload: PersonalResumeIn) -> dict[str, Any]:
    skill_cells = [x.strip() for x in re.split(r'[,;\n]', payload.skills) if x.strip()]
    achievement_cells = [x.strip() for x in re.split(r'\n|;', payload.achievements) if x.strip()]
    education = [{'degree': payload.education or 'Education details to be completed', 'school': '', 'details': ''}]
    employer = payload.employer_role or 'Most Recent Employer / Role'
    return normalize_resume_data({
        'contact': {'name': '', 'title': payload.title, 'email': '', 'phone': '', 'linkedin': '', 'portfolio': '', 'location': ''},
        'summary': f"{payload.tone.title()} professional with {payload.years or 'relevant'} years of experience targeting {payload.target or payload.title} roles. Background includes {payload.skills or 'core role-aligned skills'}, measurable execution, stakeholder support, and disciplined documentation.",
        'skills': [(skill_cells[i:i+3] + ['', '', ''])[:3] for i in range(0, max(len(skill_cells), 1), 3)] or [['Leadership', 'Operations', 'Communication']],
        'technical': {'Core Skills': payload.skills or 'Add key tools and technologies', 'Target Role': payload.target or payload.title, 'Tone': payload.tone},
        'experience': [{'title': payload.title, 'company': employer, 'location': '', 'dates': '', 'bullets': achievement_cells or ['Delivered role-aligned work improving reliability, service quality, and business outcomes.', 'Partnered with stakeholders to document requirements, resolve issues, and execute priorities.', 'Applied relevant tools and processes to complete projects accurately and efficiently.']}],
        'education': education,
        'certifications': [],
        'additional': [],
        'custom_sections': [{'title': 'Selected Projects', 'bullets': achievement_cells[:4]}] if achievement_cells else []
    })

@app.post('/api/build-from-description')
def build_from_description(payload: PersonalResumeIn):
    prompt = '''Create a complete professional resume JSON from this intake. Return ONLY JSON matching this schema:
{"contact":{"name":"","title":"","email":"","phone":"","linkedin":"","portfolio":"","location":""},"summary":"","skills":[["","",""]],"technical":{"Category":"items"},"experience":[{"title":"","company":"","location":"","dates":"","bullets":[""]}],"education":[{"degree":"","school":"","details":""}],"certifications":[],"additional":[],"custom_sections":[{"title":"","bullets":[""]}]}
Make it specific, premium, truthful to the intake, and optimized for the target role.''' + '\nINTAKE:\n' + payload.model_dump_json()
    generated = claude(prompt)
    parsed = extract_json_object(generated or '')
    resume_data = normalize_resume_data(parsed) if parsed else fallback_personal_resume(payload)
    return save_resume(ResumeIn(name=f"{payload.title or 'AI'} Resume Draft", title=payload.title or 'Resume Draft', template=payload.template, data=resume_data))

def render_text(data):
    data = merge_resume_data(data)
    c = data['contact']
    lines = [c['name'], c['title'], f"{c['email']} | {c['phone']} | {c['location']}", c.get('linkedin',''), c.get('portfolio',''), '', 'PROFESSIONAL SUMMARY', data.get('summary',''), '', 'AREAS OF EXPERTISE']
    lines += [' | '.join(row) for row in data.get('skills', [])]
    lines += ['', 'TECHNICAL PROFICIENCIES'] + [f'{k}: {v}' for k,v in data.get('technical', {}).items()]
    lines += ['', 'CAREER EXPERIENCE']
    for job in data.get('experience', []):
        lines.append(f"{job['title']} | {job['company']} | {job['location']} | {job['dates']}")
        lines += ['• ' + b for b in job.get('bullets', [])]
    lines += ['', 'EDUCATION'] + [f"{e.get('degree')} — {e.get('school')} ({e.get('details','')})" for e in data.get('education', [])]
    lines += ['', 'CERTIFICATIONS'] + data.get('certifications', [])
    lines += ['', 'ADDITIONAL EXPERIENCE'] + data.get('additional', [])
    return '\n'.join(lines)

def resume_export_response(fmt: str, data: dict[str, Any], template: str, name: str = 'resume', disposition: str = 'attachment'):
    template = template or 'modern'
    merged = merge_resume_data(data)
    filename = safe_filename(merged.get('contact', {}).get('name') or name, template, fmt)
    if fmt == 'docx':
        return Response(
            docx_bytes(merged, template),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'{disposition}; filename="{filename}"'}
        )
    if fmt == 'pdf':
        return Response(
            pdf_bytes(merged, template),
            media_type='application/pdf',
            headers={'Content-Disposition': f'{disposition}; filename="{filename}"', 'Cache-Control': 'no-store'}
        )
    raise HTTPException(400, 'fmt must be docx or pdf')

@app.get('/api/resumes/{resume_id}/export/{fmt}')
def export_resume(resume_id: int, fmt: str):
    r = resume(resume_id)
    return resume_export_response(fmt, r['data'], r.get('template') or 'modern', r.get('name') or 'resume')

@app.post('/api/resumes/export/{fmt}')
def export_resume_payload(fmt: str, payload: ResumeRenderIn):
    return resume_export_response(fmt, payload.data, payload.template, payload.name)

@app.post('/api/resume-preview/pdf')
def preview_resume_pdf(payload: ResumeRenderIn):
    return resume_export_response('pdf', payload.data, payload.template, payload.name, disposition='inline')

@app.post('/api/documents/export/{fmt}')
def export_document(fmt: str, payload: DocIn):
    return letter_export_response(fmt, payload.title, payload.data.get('content', ''))

@app.get('/api/documents/{doc_id}/export/{fmt}')
def export_document_by_id(doc_id: int, fmt: str):
    d = document(doc_id)
    return letter_export_response(fmt, str(d['title']), d['data'].get('content', ''))

def letter_export_response(fmt: str, title: str, content: str):
    safe = re.sub(r'[^A-Za-z0-9_.-]+', '-', title).strip('-').lower() or 'letter'
    if fmt == 'docx':
        doc = Document()
        for i, para in enumerate(content.split('\n')):
            p = doc.add_paragraph(para)
            if i == 0:
                for run in p.runs: run.bold = True
        buf = io.BytesIO(); doc.save(buf)
        return Response(buf.getvalue(), media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={'Content-Disposition': f'attachment; filename={safe}.docx'})
    if fmt == 'pdf':
        buf = io.BytesIO(); pdf = SimpleDocTemplate(buf, pagesize=LETTER, leftMargin=.75*inch, rightMargin=.75*inch, topMargin=.65*inch, bottomMargin=.65*inch); story=[]; styles=getSampleStyleSheet()
        for para in content.split('\n'):
            if para.strip(): story.append(Paragraph(para.replace('&','&amp;'), styles['BodyText']))
            story.append(Spacer(1,8))
        pdf.build(story)
        return Response(buf.getvalue(), media_type='application/pdf', headers={'Content-Disposition': f'attachment; filename={safe}.pdf'})
    raise HTTPException(400, 'fmt must be docx or pdf')

static_dir = Path(os.getenv('FRONTEND_DIST_DIR', '/app/frontend-dist'))
if static_dir.exists():
    app.mount('/', StaticFiles(directory=static_dir, html=True), name='static')
