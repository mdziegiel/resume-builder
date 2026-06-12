import html
import io
import re
from copy import deepcopy
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import KeepInFrame, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

PX_TO_PT = 0.75
TECHNICAL_SIDEBAR_PX = 200
TECHNICAL_SIDEBAR_WIDTH_PT = TECHNICAL_SIDEBAR_PX * PX_TO_PT
TECHNICAL_CHARCOAL = '1e293b'

from seed import DEFAULT_RESUME

TEMPLATE_IDS = {
    'classic', 'modern', 'executive', 'technical', 'minimal', 'two-column', 'corporate', 'ats-optimized'
}

ACCENT_ORANGE = 'f97316'
NAVY = '071832'
GOLD = 'c8a45d'
DARK = '111827'
GRAY = 'd1d5db'
LIGHT = 'f8fafc'

SECTION_DEFAULT_TITLES = {
    'contact': 'Contact Info',
    'summary': 'Professional Summary',
    'skills': 'Areas of Expertise',
    'technical': 'Technical Proficiencies',
    'experience': 'Career Experience',
    'education': 'Education',
    'certifications': 'Certifications',
    'additional': 'Additional Experience',
}


def section_meta(src: dict[str, Any] | None, key: str) -> dict[str, Any]:
    meta = (src or {}).get('section_meta') or {}
    item = meta.get(key) if isinstance(meta, dict) else {}
    if not isinstance(item, dict):
        item = {}
    title = strip_html(item.get('title')) or SECTION_DEFAULT_TITLES.get(key, key.replace('_', ' ').title())
    return {'title': title, 'hidden': bool(item.get('hidden'))}


def apply_section_state(out: dict[str, Any], src: dict[str, Any] | None) -> dict[str, Any]:
    out['section_meta'] = {key: section_meta(src, key) for key in SECTION_DEFAULT_TITLES}
    for key in SECTION_DEFAULT_TITLES:
        if out['section_meta'][key]['hidden']:
            if key == 'contact':
                out['contact'] = {k: '' for k in out.get('contact', {})}
            elif key == 'summary':
                out['summary'] = ''
            elif key == 'technical':
                out['technical'] = {}
            elif key == 'skills':
                out['skills'] = []
            elif key == 'experience':
                out['experience'] = []
            elif key == 'education':
                out['education'] = []
            elif key == 'certifications':
                out['certifications'] = []
            elif key == 'additional':
                out['additional'] = []
    return out


def section_visible(data: dict[str, Any], key: str) -> bool:
    return not ((data.get('section_meta') or {}).get(key) or {}).get('hidden')


def section_title(data: dict[str, Any], key: str) -> str:
    return ((data.get('section_meta') or {}).get(key) or {}).get('title') or SECTION_DEFAULT_TITLES.get(key, key.replace('_', ' ').title())


def strip_html(value: Any) -> str:
    text = re.sub(r'<\s*br\s*/?>', '\n', str(value or ''), flags=re.I)
    text = re.sub(r'</\s*(p|div|li)\s*>', '\n', text, flags=re.I)
    text = re.sub(r'<[^>]+>', '', text)
    return html.unescape(text).replace('\xa0', ' ').strip()


def has_text(value: Any) -> bool:
    return bool(strip_html(value))


def merge_resume_data(data: dict[str, Any] | None) -> dict[str, Any]:
    """Layer real fields over fictional preview defaults, field-by-field."""
    src = deepcopy(data or {})
    base = deepcopy(DEFAULT_RESUME)
    out = deepcopy(base)
    contact = src.get('contact') or {}
    out['contact'] = {k: strip_html(contact.get(k)) or base['contact'].get(k, '') for k in base['contact']}
    out['summary'] = strip_html(src.get('summary')) or base['summary']

    skills = normalize_skill_rows(src.get('skills'))
    out['skills'] = skills if any(any(has_text(c) for c in r) for r in skills) else base['skills']

    technical = {strip_html(k): strip_html(v) for k, v in (src.get('technical') or {}).items() if strip_html(k) and strip_html(v)}
    out['technical'] = technical or base['technical']

    experience = []
    for item in src.get('experience') or []:
        bullets = [strip_html(b) for b in item.get('bullets', []) if has_text(b)]
        if any(has_text(item.get(k)) for k in ('title', 'company', 'location', 'dates', 'start_date', 'end_date')) or bullets:
            dates = strip_html(item.get('dates')) or ' – '.join(x for x in [strip_html(item.get('start_date')), strip_html(item.get('end_date'))] if x)
            experience.append({
                'title': strip_html(item.get('title')),
                'company': strip_html(item.get('company')),
                'location': strip_html(item.get('location')),
                'dates': dates,
                'bullets': bullets,
            })
    out['experience'] = experience or base['experience']

    education = []
    for item in src.get('education') or []:
        if any(has_text(item.get(k)) for k in ('degree', 'school', 'details')):
            education.append({k: strip_html(item.get(k)) for k in ('degree', 'school', 'details')})
    out['education'] = education or base['education']

    certs = [strip_html(x) for x in src.get('certifications', []) if has_text(x)]
    out['certifications'] = certs or base['certifications']
    additional = [strip_html(x) for x in src.get('additional', []) if has_text(x)]
    out['additional'] = additional or base['additional']
    custom = []
    for section in src.get('custom_sections') or []:
        title = strip_html(section.get('title'))
        bullets = [strip_html(x) for x in section.get('bullets', []) if has_text(x)]
        if title or bullets:
            custom.append({'title': title or 'Additional Information', 'bullets': bullets})
    out['custom_sections'] = custom
    return apply_section_state(out, src)


def normalize_skill_rows(skills: Any) -> list[list[str]]:
    if isinstance(skills, dict):
        source = list(skills.values())
    elif isinstance(skills, list):
        source = skills
    elif skills:
        source = [skills]
    else:
        source = []
    cells: list[str] = []
    for row in source:
        if isinstance(row, dict):
            items = row.get('items')
            if isinstance(items, list):
                vals = items
            elif isinstance(items, str):
                vals = re.split(r'[|,;•\n]', items)
            else:
                vals = [row.get('a'), row.get('b'), row.get('c'), row.get('name'), row.get('label'), row.get('value')]
        elif isinstance(row, str):
            vals = re.split(r'[|,;•\n]', row)
        elif isinstance(row, (list, tuple)):
            vals = list(row)
        else:
            vals = []
        cells.extend(strip_html(x) for x in vals if has_text(x))
    return [(cells[i:i + 3] + ['', '', ''])[:3] for i in range(0, len(cells), 3)]


def safe_filename(name: str, template: str, ext: str) -> str:
    root = re.sub(r'[^A-Za-z0-9_.-]+', '-', f'{name or "resume"}-{template}').strip('-').lower()
    return f'{root or "resume"}.{ext}'


def docx_bytes(data: dict[str, Any], template: str) -> bytes:
    template = template if template in TEMPLATE_IDS else 'modern'
    data = merge_resume_data(data)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    if template == 'technical':
        # A true fixed-width sidebar needs to own the page edge and table grid.
        # Normal Word margins plus percentage widths are what caused overlap in exports.
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    else:
        section.top_margin = Inches(0.48)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.58)
        section.right_margin = Inches(0.58)
    set_doc_defaults(doc, template)

    if template == 'technical':
        render_docx_sidebar(doc, data, template, dark=True, width_pct=32)
    elif template == 'two-column':
        render_docx_sidebar(doc, data, template, dark=False, width_pct=30)
    elif template == 'ats-optimized':
        render_docx_ats(doc, data)
    else:
        render_docx_single(doc, data, template)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def set_doc_defaults(doc: Document, template: str) -> None:
    font = {
        'classic': 'Georgia', 'modern': 'Calibri', 'executive': 'Garamond', 'technical': 'Calibri',
        'minimal': 'Calibri Light', 'two-column': 'Calibri', 'corporate': 'Calibri', 'ats-optimized': 'Arial'
    }.get(template, 'Calibri')
    styles = doc.styles
    styles['Normal'].font.name = font
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    styles['Normal'].font.size = Pt(9.5 if template != 'ats-optimized' else 10)
    styles['Normal'].paragraph_format.space_after = Pt(3)
    for style_name in ('Heading 1', 'Heading 2'):
        styles[style_name].font.name = font
        styles[style_name]._element.rPr.rFonts.set(qn('w:eastAsia'), font)
        styles[style_name].font.bold = True
        styles[style_name].font.size = Pt(11)
        styles[style_name].paragraph_format.space_before = Pt(10)
        styles[style_name].paragraph_format.space_after = Pt(4)


def rgb(hex_color: str) -> RGBColor:
    h = hex_color.strip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_cell_width(cell, width_twips: int) -> None:
    cell.width = Inches(width_twips / 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_twips))
    tc_w.set(qn('w:type'), 'dxa')


def set_fixed_table_layout(table) -> None:
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn('w:tblLayout'))
    if tbl_layout is None:
        tbl_layout = OxmlElement('w:tblLayout')
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn('w:type'), 'fixed')


def para_border(paragraph, bottom: str | None = None, left: str | None = None, top: str | None = None, size='8', space='2') -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn('w:pBdr'))
    if borders is None:
        borders = OxmlElement('w:pBdr')
        p_pr.append(borders)
    for side, color in [('top', top), ('bottom', bottom), ('left', left)]:
        if not color:
            continue
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), size)
        node.set(qn('w:space'), space)
        node.set(qn('w:color'), color)
        borders.append(node)


def add_run(paragraph, text: str, *, bold=False, color: str | None = None, size: float | None = None, font: str | None = None, caps=False):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = rgb(color)
    if size:
        run.font.size = Pt(size)
    if font:
        run.font.name = font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if caps:
        run.font.small_caps = True
    return run


def section_heading(doc_or_cell, text: str, template: str):
    p = doc_or_cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(9 if template != 'minimal' else 15)
    pf.space_after = Pt(4)
    if template == 'classic':
        add_run(p, text, bold=True, size=11, font='Georgia', caps=True)
        para_border(p, bottom=DARK, size='6')
    elif template == 'modern':
        pf.left_indent = Pt(8)
        add_run(p, text.upper(), bold=True, size=11, font='Calibri')
        para_border(p, left=ACCENT_ORANGE, size='14')
    elif template == 'executive':
        add_run(p, text.upper(), bold=True, size=11, font='Garamond', color=NAVY)
        para_border(p, bottom=GOLD, size='12')
    elif template == 'minimal':
        add_run(p, text.upper(), bold=False, size=10.5, font='Calibri Light')
        para_border(p, top=GRAY, size='4')
    elif template == 'corporate':
        add_run(p, text.upper(), bold=True, size=11, font='Calibri', color=NAVY)
        para_border(p, bottom=NAVY, size='8')
    elif template in ('technical', 'two-column'):
        add_run(p, text.upper(), bold=True, size=10.5, font='Calibri', color='0f172a')
        para_border(p, bottom='cbd5e1', size='6')
    else:
        add_run(p, text.upper(), bold=True, size=10.5, font='Arial')
    return p


def render_docx_single(doc: Document, data: dict[str, Any], template: str) -> None:
    c = data['contact']
    if template in ('executive', 'corporate'):
        table = doc.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.cell(0, 0)
        shade_cell(cell, NAVY if template == 'executive' else '0f172a')
        cell_margins(cell, 360, 420, 320, 420)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if template == 'corporate' else WD_ALIGN_PARAGRAPH.LEFT
        name = spaced(c['name'].upper()) if template == 'executive' else c['name']
        add_run(p, name, bold=True, color='ffffff', size=24 if template == 'executive' else 23, font='Garamond' if template == 'executive' else 'Calibri')
        p2 = cell.add_paragraph()
        p2.alignment = p.alignment
        add_run(p2, c['title'], color='e5e7eb', size=11, font='Garamond' if template == 'executive' else 'Calibri')
        p3 = cell.add_paragraph()
        p3.alignment = p.alignment
        add_run(p3, contact_line(c), color='e5e7eb', size=8.5, font='Garamond' if template == 'executive' else 'Calibri')
        if template == 'corporate':
            p4 = doc.add_paragraph()
            para_border(p4, bottom=GOLD, size='18')
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if template == 'classic' else WD_ALIGN_PARAGRAPH.LEFT
        font = {'classic': 'Georgia', 'modern': 'Calibri', 'minimal': 'Calibri Light'}.get(template, 'Calibri')
        size = {'classic': 18, 'modern': 22, 'minimal': 24}.get(template, 22)
        add_run(p, c['name'], bold=(template != 'minimal'), size=size, font=font, color='111111')
        p.paragraph_format.space_after = Pt(2)
        p2 = doc.add_paragraph()
        p2.alignment = p.alignment
        add_run(p2, c['title'], size=10.5, font=font, color=ACCENT_ORANGE if template == 'modern' else '374151')
        p3 = doc.add_paragraph()
        p3.alignment = p.alignment
        add_run(p3, contact_line(c), size=8.8, font=font, color='374151')
        if template == 'classic':
            para_border(p3, bottom=DARK, top=DARK, size='6')
        elif template == 'modern':
            para_border(p3, bottom=ACCENT_ORANGE, size='16')
        elif template == 'minimal':
            para_border(p3, bottom=GRAY, size='4')

    render_docx_body(doc, data, template, corporate_alt=(template == 'corporate'))


def render_docx_body(container, data: dict[str, Any], template: str, *, corporate_alt=False) -> None:
    if section_visible(data, 'summary') and data.get('summary'):
        section_heading(container, section_title(data, 'summary'), template)
        container.add_paragraph(data['summary'])
    if section_visible(data, 'skills') and data.get('skills'):
        section_heading(container, section_title(data, 'skills'), template)
        add_skills_table(container, data['skills'], template)
    if section_visible(data, 'technical') and data.get('technical'):
        section_heading(container, section_title(data, 'technical'), template)
        for k, v in data['technical'].items():
            p = container.add_paragraph()
            add_run(p, f'{k}: ', bold=True, font='Consolas' if template == 'technical' else None)
            add_run(p, v, font='Consolas' if template == 'technical' else None)
    if section_visible(data, 'experience') and data.get('experience'):
        section_heading(container, section_title(data, 'experience'), template)
        for idx, job in enumerate(data['experience']):
            if corporate_alt and idx % 2 == 1:
                t = container.add_table(rows=1, cols=1)
                cell = t.cell(0, 0)
                shade_cell(cell, 'f9fafb')
                cell_margins(cell, 90, 120, 80, 120)
                render_job_docx(cell, job, template)
            else:
                render_job_docx(container, job, template)
    if section_visible(data, 'education') and data.get('education'):
        section_heading(container, section_title(data, 'education'), template)
        for e in data['education']:
            container.add_paragraph(' — '.join(x for x in [e.get('degree'), e.get('school'), e.get('details')] if x))
    if section_visible(data, 'certifications') and data.get('certifications'):
        section_heading(container, section_title(data, 'certifications'), template)
        container.add_paragraph(' | '.join(data['certifications']))
    if section_visible(data, 'additional') and data.get('additional'):
        section_heading(container, section_title(data, 'additional'), template)
        for item in data['additional']:
            container.add_paragraph(item)
    for custom in data.get('custom_sections', []):
        if custom.get('hidden'):
            continue
        section_heading(container, custom['title'], template)
        for bullet in custom.get('bullets', []):
            container.add_paragraph(bullet, style='List Bullet')


def render_job_docx(container, job: dict[str, Any], template: str) -> None:
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    add_run(p, ' — '.join(x for x in [job.get('title'), job.get('company')] if x), bold=True, size=9.7)
    if job.get('location') or job.get('dates'):
        p2 = container.add_paragraph()
        add_run(p2, ' | '.join(x for x in [job.get('location'), job.get('dates')] if x), color='4b5563', size=8.8)
    for bullet in job.get('bullets', []):
        container.add_paragraph(bullet, style='List Bullet')


def add_skills_table(container, skills: list[list[str]], template: str) -> None:
    if template == 'ats-optimized':
        container.add_paragraph(' | '.join([c for row in skills for c in row if c]))
        return
    table = container.add_table(rows=len(skills), cols=3)
    table.autofit = True
    for i, row in enumerate(skills):
        for j in range(3):
            cell = table.cell(i, j)
            cell.text = row[j] if j < len(row) else ''
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.2)
                p.paragraph_format.keep_together = True
            tc_pr = cell._tc.get_or_add_tcPr()
            no_wrap = OxmlElement('w:noWrap')
            tc_pr.append(no_wrap)
            cell_margins(cell, 34, 42, 34, 42)
            if template == 'two-column':
                shade_cell(cell, 'e2e8f0')
            elif template == 'technical':
                shade_cell(cell, '334155')
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = rgb('ffffff')
            elif template == 'minimal':
                shade_cell(cell, 'ffffff')


def render_docx_sidebar(doc: Document, data: dict[str, Any], template: str, *, dark: bool, width_pct: int) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_fixed_table_layout(table)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    if template == 'technical':
        page_twips = int(doc.sections[0].page_width.inches * 1440)
        left_twips = int((TECHNICAL_SIDEBAR_WIDTH_PT / 72) * 1440)
        right_twips = page_twips - left_twips
        sidebar_fill = TECHNICAL_CHARCOAL
        left_margins = (620, 190, 620, 190)
        right_margins = (620, 420, 620, 420)
    else:
        usable_twips = int((doc.sections[0].page_width.inches - doc.sections[0].left_margin.inches - doc.sections[0].right_margin.inches) * 1440)
        left_twips = int(usable_twips * width_pct / 100)
        right_twips = usable_twips - left_twips
        sidebar_fill = 'f8fafc'
        left_margins = (360, 220, 360, 220)
        right_margins = (360, 340, 360, 340)
    set_cell_width(left, left_twips)
    set_cell_width(right, right_twips)
    shade_cell(left, sidebar_fill)
    shade_cell(right, 'ffffff')
    cell_margins(left, *left_margins)
    cell_margins(right, *right_margins)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    c = data['contact']
    p = left.paragraphs[0]
    add_run(p, c['name'], bold=True, color='ffffff' if dark else '0f172a', size=18, font='Calibri')
    p2 = left.add_paragraph()
    add_run(p2, c['title'], color='93c5fd' if dark else '475569', size=9.5)
    sidebar_heading(left, 'Contact', dark)
    for item in [c.get('email'), c.get('phone'), c.get('location'), c.get('linkedin'), c.get('portfolio')]:
        if item:
            p = left.add_paragraph()
            add_run(p, item, color='ffffff' if dark else '334155', size=8.2)
    sidebar_heading(left, 'Skills', dark)
    for skill in [x for row in data['skills'] for x in row if x]:
        p = left.add_paragraph()
        add_run(p, skill, color='ffffff' if dark else '0f172a', size=8.4, font='Consolas' if template == 'technical' else 'Calibri')
    sidebar_heading(left, 'Certifications', dark)
    for cert in data['certifications']:
        p = left.add_paragraph()
        add_run(p, cert, color='ffffff' if dark else '334155', size=8.2)
    render_docx_body(right, data, template)


def sidebar_heading(container, text: str, dark: bool) -> None:
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text.upper(), bold=True, color='93c5fd' if dark else '0f172a', size=9)
    para_border(p, bottom='ffffff' if dark else 'cbd5e1', size='4')


def render_docx_ats(doc: Document, data: dict[str, Any]) -> None:
    c = data['contact']
    for text, bold in [(c['name'], True), (c['title'], False), (contact_line(c), False)]:
        p = doc.add_paragraph()
        add_run(p, text, bold=bold, size=10.5, font='Arial')
    for heading, body in [('Professional Summary', [data['summary']]), ('Areas of Expertise', [' | '.join(x for r in data['skills'] for x in r if x)]), ('Technical Proficiencies', [f'{k}: {v}' for k, v in data['technical'].items()])]:
        section_heading(doc, heading, 'ats-optimized')
        for line in body:
            doc.add_paragraph(line)
    section_heading(doc, 'Career Experience', 'ats-optimized')
    for job in data['experience']:
        doc.add_paragraph(' | '.join(x for x in [job.get('title'), job.get('company'), job.get('location'), job.get('dates')] if x))
        for bullet in job.get('bullets', []):
            doc.add_paragraph(bullet)
    section_heading(doc, 'Education', 'ats-optimized')
    for e in data['education']:
        doc.add_paragraph(' | '.join(x for x in [e.get('degree'), e.get('school'), e.get('details')] if x))
    section_heading(doc, 'Certifications', 'ats-optimized')
    doc.add_paragraph(' | '.join(data['certifications']))


def contact_line(c: dict[str, Any]) -> str:
    return ' | '.join(x for x in [c.get('email'), c.get('phone'), c.get('location'), c.get('linkedin'), c.get('portfolio')] if x)


def spaced(text: str) -> str:
    return ' '.join(list(text))


def pdf_bytes(data: dict[str, Any], template: str) -> bytes:
    template = template if template in TEMPLATE_IDS else 'modern'
    data = merge_resume_data(data)
    buf = io.BytesIO()
    margins = 0.42 * inch
    top_margin = 0.45 * inch
    bottom_margin = 0.45 * inch
    if template == 'minimal':
        margins = 0.50 * inch
    elif template == 'technical':
        margins = 0
        top_margin = 0
        bottom_margin = 0
    doc = SimpleDocTemplate(buf, pagesize=LETTER, leftMargin=margins, rightMargin=margins, topMargin=top_margin, bottomMargin=bottom_margin)
    styles = pdf_styles(template)
    story = []
    if template == 'technical':
        story.extend(pdf_sidebar(data, template, dark=True))
    elif template == 'two-column':
        story.extend(pdf_sidebar(data, template, dark=False))
    elif template == 'ats-optimized':
        story.extend(pdf_ats(data, styles))
    else:
        story.extend(pdf_single(data, template, styles))
    doc.build(story)
    return buf.getvalue()


def pdf_styles(template: str) -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    font = {
        'classic': 'Times-Roman', 'executive': 'Times-Roman', 'technical': 'Helvetica', 'minimal': 'Helvetica',
        'two-column': 'Helvetica', 'corporate': 'Helvetica', 'ats-optimized': 'Helvetica', 'modern': 'Helvetica'
    }.get(template, 'Helvetica')
    return {
        'body': ParagraphStyle('body', parent=base['BodyText'], fontName=font, fontSize=8.8 if template != 'ats-optimized' else 9.3, leading=11.2, spaceAfter=3),
        'name': ParagraphStyle('name', parent=base['Title'], fontName='Times-Bold' if template in ('classic', 'executive') else 'Helvetica-Bold', fontSize={'classic': 18, 'modern': 22, 'executive': 24, 'minimal': 24, 'corporate': 23}.get(template, 20), leading=25, spaceAfter=2),
        'title': ParagraphStyle('title', parent=base['BodyText'], fontName=font, fontSize=10.2, leading=12, spaceAfter=4, textColor=colors.HexColor('#374151')),
        'section': ParagraphStyle('section', parent=base['Heading2'], fontName='Helvetica-Bold', fontSize=10, leading=12, spaceBefore=8, spaceAfter=5, textTransform='uppercase', textColor=colors.HexColor('#111827')),
        'small': ParagraphStyle('small', parent=base['BodyText'], fontName=font, fontSize=8, leading=9.8, textColor=colors.HexColor('#4b5563')),
    }


def P(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(strip_html(text)).replace('\n', '<br/>'), style)


def pdf_heading(text: str, template: str, styles: dict[str, ParagraphStyle]):
    color = {'modern': '#f97316', 'executive': '#071832', 'corporate': '#071832'}.get(template, '#111827')
    bg = '#fff7ed' if template == 'modern' else None
    table = Table([[P(text.upper(), ParagraphStyle('h', parent=styles['section'], textColor=colors.HexColor(color)))]], colWidths=['100%'])
    commands = [('BOTTOMPADDING', (0, 0), (-1, -1), 3)]
    if template == 'classic':
        commands += [('LINEBELOW', (0, 0), (-1, -1), 0.6, colors.HexColor('#111827'))]
    elif template == 'modern':
        commands += [('LINEBEFORE', (0, 0), (0, 0), 3, colors.HexColor('#f97316')), ('LEFTPADDING', (0, 0), (0, 0), 8)]
        if bg:
            commands += [('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(bg))]
    elif template == 'executive':
        commands += [('LINEBELOW', (0, 0), (-1, -1), 1.5, colors.HexColor('#c8a45d'))]
    elif template == 'minimal':
        commands += [('LINEABOVE', (0, 0), (-1, -1), 0.4, colors.HexColor('#d1d5db'))]
    elif template == 'corporate':
        commands += [('LINEBELOW', (0, 0), (-1, -1), 1, colors.HexColor('#071832'))]
    table.setStyle(TableStyle(commands))
    return [Spacer(1, 5), table, Spacer(1, 2)]


def pdf_single(data: dict[str, Any], template: str, styles: dict[str, ParagraphStyle]) -> list[Any]:
    c = data['contact']
    story: list[Any] = []
    if template in ('executive', 'corporate'):
        name_style = ParagraphStyle('blockname', parent=styles['name'], alignment=TA_CENTER if template == 'corporate' else TA_LEFT, textColor=colors.white, fontName='Times-Bold' if template == 'executive' else 'Helvetica-Bold')
        title_style = ParagraphStyle('blocktitle', parent=styles['title'], alignment=name_style.alignment, textColor=colors.HexColor('#e5e7eb'))
        header = [[P(spaced(c['name'].upper()) if template == 'executive' else c['name'], name_style)], [P(c['title'], title_style)], [P(contact_line(c), title_style)]]
        t = Table(header, colWidths=[7.4 * inch])
        t.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#071832')), ('BOX', (0, 0), (-1, -1), 0, colors.HexColor('#071832')), ('TOPPADDING', (0, 0), (-1, -1), 10), ('BOTTOMPADDING', (0, 0), (-1, -1), 8), ('LEFTPADDING', (0, 0), (-1, -1), 18), ('RIGHTPADDING', (0, 0), (-1, -1), 18)]))
        story += [t, Spacer(1, 12)]
    else:
        align = TA_CENTER if template == 'classic' else TA_LEFT
        name_style = ParagraphStyle('name2', parent=styles['name'], alignment=align, fontName='Times-Bold' if template == 'classic' else 'Helvetica', textColor=colors.black)
        if template in ('modern',):
            name_style.fontName = 'Helvetica-Bold'
        title_style = ParagraphStyle('title2', parent=styles['title'], alignment=align, textColor=colors.HexColor('#f97316') if template == 'modern' else colors.HexColor('#374151'))
        contact_style = ParagraphStyle('contact', parent=styles['small'], alignment=align)
        story += [P(c['name'], name_style), P(c['title'], title_style), P(contact_line(c), contact_style)]
        rule_color = {'classic': '#111827', 'modern': '#f97316', 'minimal': '#d1d5db'}.get(template, '#d1d5db')
        width = 1.2 if template == 'modern' else 0.5
        rule = Table([['']], colWidths=[7.4 * inch], rowHeights=[2])
        rule.setStyle(TableStyle([('LINEBELOW', (0, 0), (-1, -1), width, colors.HexColor(rule_color))]))
        story += [rule, Spacer(1, 7)]
    story += pdf_body(data, template, styles)
    return story


def pdf_body(data: dict[str, Any], template: str, styles: dict[str, ParagraphStyle]) -> list[Any]:
    story: list[Any] = []
    if section_visible(data, 'summary') and data.get('summary'):
        story += pdf_heading(section_title(data, 'summary'), template, styles)
        story.append(P(data['summary'], styles['body']))
    if section_visible(data, 'skills') and data.get('skills'):
        story += pdf_heading(section_title(data, 'skills'), template, styles)
        skill_line = ' | '.join(x for row in data['skills'] for x in row if x)
        if template == 'ats-optimized':
            story.append(P(skill_line, styles['body']))
        elif template in ('two-column',):
            story.append(chip_table([x for row in data['skills'] for x in row if x], '#e2e8f0'))
        else:
            skill_style = ParagraphStyle('skillwrap', parent=styles['body'], fontSize=7.4, leading=8.8, splitLongWords=True)
            rows = [[P(c, skill_style) for c in row] for row in data['skills']]
            t = Table(rows, colWidths=[2.48 * inch, 2.48 * inch, 2.48 * inch])
            t.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP'), ('LEFTPADDING', (0, 0), (-1, -1), 0), ('RIGHTPADDING', (0, 0), (-1, -1), 2), ('BOTTOMPADDING', (0, 0), (-1, -1), 1)]))
            story.append(t)
    if section_visible(data, 'technical') and data.get('technical'):
        story += pdf_heading(section_title(data, 'technical'), template, styles)
        tech_style = ParagraphStyle('tech', parent=styles['body'], fontName='Courier' if template == 'technical' else styles['body'].fontName)
        for k, v in data['technical'].items():
            story.append(Paragraph(f'<b>{html.escape(k)}:</b> {html.escape(v)}', tech_style))
    if section_visible(data, 'experience') and data.get('experience'):
        story += pdf_heading(section_title(data, 'experience'), template, styles)
        for idx, job in enumerate(data['experience']):
            block = pdf_job(job, styles)
            if template == 'corporate' and idx % 2 == 1:
                t = Table([[block]], colWidths=[7.25 * inch])
                t.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f9fafb')), ('LEFTPADDING', (0, 0), (-1, -1), 8), ('RIGHTPADDING', (0, 0), (-1, -1), 8), ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 3)]))
                story.append(t)
            else:
                story.extend(block)
    if section_visible(data, 'education') and data.get('education'):
        story += pdf_heading(section_title(data, 'education'), template, styles)
        for e in data['education']:
            story.append(P(' — '.join(x for x in [e.get('degree'), e.get('school'), e.get('details')] if x), styles['body']))
    if section_visible(data, 'certifications') and data.get('certifications'):
        story += pdf_heading(section_title(data, 'certifications'), template, styles)
        story.append(P(' | '.join(data['certifications']), styles['body']))
    if section_visible(data, 'additional') and data.get('additional'):
        story += pdf_heading(section_title(data, 'additional'), template, styles)
        for a in data['additional']:
            story.append(P(a, styles['body']))
    for custom in data.get('custom_sections', []):
        if custom.get('hidden'):
            continue
        story += pdf_heading(custom['title'], template, styles)
        for b in custom.get('bullets', []):
            story.append(P('• ' + b, styles['body']))
    return story


def pdf_job(job: dict[str, Any], styles: dict[str, ParagraphStyle]) -> list[Any]:
    title = ' — '.join(x for x in [job.get('title'), job.get('company')] if x)
    meta = ' | '.join(x for x in [job.get('location'), job.get('dates')] if x)
    out = [Paragraph(f'<b>{html.escape(title)}</b>', styles['body'])]
    if meta:
        out.append(P(meta, styles['small']))
    for b in job.get('bullets', []):
        out.append(P('• ' + b, styles['body']))
    out.append(Spacer(1, 4))
    return out


def chip_table(items: list[str], fill: str) -> Table:
    rows = []
    for i in range(0, len(items), 2):
        rows.append([Paragraph(html.escape(items[i]), getSampleStyleSheet()['BodyText']), Paragraph(html.escape(items[i+1]), getSampleStyleSheet()['BodyText']) if i + 1 < len(items) else ''])
    t = Table(rows, colWidths=[1.8 * inch, 1.8 * inch])
    t.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(fill)), ('BOX', (0, 0), (-1, -1), 0.2, colors.HexColor('#cbd5e1')), ('INNERGRID', (0, 0), (-1, -1), 3, colors.white), ('LEFTPADDING', (0, 0), (-1, -1), 6), ('RIGHTPADDING', (0, 0), (-1, -1), 6), ('TOPPADDING', (0, 0), (-1, -1), 3), ('BOTTOMPADDING', (0, 0), (-1, -1), 3)]))
    return t


def pdf_sidebar(data: dict[str, Any], template: str, dark: bool) -> list[Any]:
    styles = pdf_styles(template)
    c = data['contact']
    if template == 'technical':
        page_width, page_height = LETTER[0] - 12, LETTER[1] - 12
        left_col = TECHNICAL_SIDEBAR_WIDTH_PT
        right_col = page_width - left_col
        row_height = page_height
        side_pad_l, side_pad_r = 13, 11
        main_pad_l, main_pad_r = 28, 28
        side_frame_width = left_col - side_pad_l - side_pad_r
        main_frame_width = right_col - main_pad_l - main_pad_r
        frame_height = row_height - 56
        sidebar_style = ParagraphStyle('side', parent=styles['small'], textColor=colors.white, fontName='Helvetica', fontSize=7.2, leading=8.8, splitLongWords=True, wordWrap='CJK')
        side_name_style = ParagraphStyle('sidename', parent=sidebar_style, fontName='Helvetica-Bold', fontSize=13.5, leading=15.5, splitLongWords=True, wordWrap='CJK')
        side_head_style = ParagraphStyle('sideh', parent=sidebar_style, fontName='Helvetica-Bold', fontSize=7.6, leading=9.2, textColor=colors.HexColor('#93c5fd'))
        main_parts = pdf_body(data, template, styles)
    else:
        page_width = LETTER[0] - (0.42 * inch * 2)
        left_col = page_width * 0.30
        right_col = page_width - left_col
        row_height = 9.75 * inch
        side_pad_l, side_pad_r = 14, 12
        main_pad_l, main_pad_r = 20, 0
        side_frame_width = left_col - side_pad_l - side_pad_r
        main_frame_width = right_col - main_pad_l - main_pad_r
        frame_height = 9.15 * inch
        sidebar_style = ParagraphStyle('side', parent=styles['small'], textColor=colors.HexColor('#0f172a'), fontName='Helvetica', fontSize=7.8, leading=9.5, splitLongWords=True)
        side_name_style = ParagraphStyle('sidename', parent=sidebar_style, fontName='Helvetica-Bold', fontSize=15, leading=17)
        side_head_style = ParagraphStyle('sideh', parent=sidebar_style, fontName='Helvetica-Bold', textColor=colors.HexColor('#0f172a'))
        main_parts = pdf_body(data, template, styles)
    side_parts: list[Any] = [P(c['name'], side_name_style), P(c['title'], sidebar_style), Spacer(1, 8)]
    for head, values in [('Contact', [c.get('email'), c.get('phone'), c.get('location'), c.get('linkedin'), c.get('portfolio')]), ('Skills', [x for row in data['skills'] for x in row if x]), ('Certifications', data['certifications'])]:
        side_parts.append(P(head.upper(), side_head_style))
        for v in values:
            if v:
                side_parts.append(P(v, sidebar_style))
        side_parts.append(Spacer(1, 6))
    side_frame = KeepInFrame(side_frame_width, frame_height, side_parts, mode='shrink')
    main_frame = KeepInFrame(main_frame_width, frame_height, main_parts, mode='shrink')
    table = Table([[side_frame, main_frame]], colWidths=[left_col, right_col], rowHeights=[row_height])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, 0), colors.HexColor('#' + (TECHNICAL_CHARCOAL if dark else 'f8fafc'))),
        ('BACKGROUND', (1, 0), (1, 0), colors.white),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (0, 0), side_pad_l),
        ('RIGHTPADDING', (0, 0), (0, 0), side_pad_r),
        ('TOPPADDING', (0, 0), (-1, -1), 28 if template == 'technical' else 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 28 if template == 'technical' else 8),
        ('LEFTPADDING', (1, 0), (1, 0), main_pad_l),
        ('RIGHTPADDING', (1, 0), (1, 0), main_pad_r),
        ('LINEBEFORE', (1, 0), (1, 0), 1.0, colors.HexColor('#38bdf8' if template == 'technical' else '#cbd5e1')),
    ]))
    return [table]


def pdf_ats(data: dict[str, Any], styles: dict[str, ParagraphStyle]) -> list[Any]:
    c = data['contact']
    plain = ParagraphStyle('plain', parent=styles['body'], fontName='Helvetica', fontSize=9.5, leading=11.5, textColor=colors.black)
    story: list[Any] = [P(c['name'], ParagraphStyle('atsname', parent=plain, fontName='Helvetica-Bold', fontSize=11)), P(c['title'], plain), P(contact_line(c), plain), Spacer(1, 6)]
    sections = [
        ('summary', [data['summary']]),
        ('skills', [' | '.join(x for r in data['skills'] for x in r if x)]),
        ('technical', [f'{k}: {v}' for k, v in data['technical'].items()]),
    ]
    for key, lines in sections:
        if section_visible(data, key) and any(lines):
            story.append(P(section_title(data, key).upper(), ParagraphStyle(key, parent=plain, fontName='Helvetica-Bold')))
            for line in lines:
                story.append(P(line, plain))
    if section_visible(data, 'experience') and data.get('experience'):
        story.append(P(section_title(data, 'experience').upper(), ParagraphStyle('exp', parent=plain, fontName='Helvetica-Bold')))
        for job in data['experience']:
            story.append(P(' | '.join(x for x in [job.get('title'), job.get('company'), job.get('location'), job.get('dates')] if x), plain))
            for b in job.get('bullets', []):
                story.append(P(b, plain))
    if section_visible(data, 'education') and data.get('education'):
        story.append(P(section_title(data, 'education').upper(), ParagraphStyle('edu', parent=plain, fontName='Helvetica-Bold')))
        for e in data['education']:
            story.append(P(' | '.join(x for x in [e.get('degree'), e.get('school'), e.get('details')] if x), plain))
    if section_visible(data, 'certifications') and data.get('certifications'):
        story.append(P(section_title(data, 'certifications').upper(), ParagraphStyle('cert', parent=plain, fontName='Helvetica-Bold')))
        story.append(P(' | '.join(data['certifications']), plain))
    for custom in data.get('custom_sections', []):
        if custom.get('hidden'):
            continue
        story.append(P(custom.get('title', 'Additional Information').upper(), ParagraphStyle(custom.get('title', 'custom'), parent=plain, fontName='Helvetica-Bold')))
        for b in custom.get('bullets', []):
            story.append(P(b, plain))
    return story
